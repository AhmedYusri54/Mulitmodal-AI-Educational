import os
import openai
from langchain.embeddings import OpenAIEmbeddings
from dotenv import load_dotenv
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import ChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
import PyPDF2
import docx



load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

class DocumentProcessor:
    def __init__(self):
        self.embeddings = None
        self.vector_store = None
        self.conversation_chain = None
        self.memory = None
        self.MODEL = "gpt-4o-mini"
    def load_models(self):
        if self.embeddings is None:
            self.embeddings = OpenAIEmbeddings()
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        return text

    def extract_text_from_python_docx(self, file_path: str):
        """Extract text from python_docx file"""
        try:
            doc = python_docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading python_docx: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")

    def process_document(self, file_path: str):
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_extension == '.python_docx':
                text = self.extract_text_from_python_docx(file_path)
            elif file_extension == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                return "", "", f"Unsupported file format: {file_extension}"

            if not text.strip():
                return "", "", "No text found in the document"

            self.processed_document_text = text

            self.create_document_vector_store(text)
            summary = self.generate_document_summary(text)
            return text, summary, "Document processed successfully!"

        except Exception as e:
            return "", "", f"Error processing document: {str(e)}"

    def create_document_vector_store(self, text: str):
        """Create FAISS vector store from document text"""
        self.load_models()
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

        chunks = text_splitter.split_text(text)
        documents = [Document(page_content=chunk) for chunk in chunks]

        self.document_vector_store = FAISS.from_documents(documents, self.embeddings)

        self.setup_document_conversation_chain()

        return self.document_vector_store

    def setup_document_conversation_chain(self):
        """Setup the conversational retrieval chain for documents"""
        if self.document_vector_store is not None:
            llm = ChatOpenAI(temperature=0.7, model_name=self.MODEL)
            self.document_memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
            retriever = self.document_vector_store.as_retriever()
            self.document_conversation_chain = ConversationalRetrievalChain.from_llm(
                llm=llm,
                retriever=retriever,
                memory=self.document_memory
            )

    def generate_document_summary(self, text: str) -> str:
        """Generate summary for document using OpenAI GPT"""
        try:
            client = openai.OpenAI()
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": """You are a helpful assistant that creates comprehensive summaries of documents.

                        IMPORTANT INSTRUCTIONS:
                        1. Analyze the language of the provided document
                        2. Write the summary in the SAME LANGUAGE as the document
                        3. If the document is in Arabic, write the summary in Arabic
                        4. If the document is in English, write the summary in English
                        5. And so on for any other language

                        Create a well-structured summary that includes:
                        - Main topic and purpose of the document
                        - Key points and sections
                        - Important details and findings
                        - Conclusions or recommendations

                        Keep the summary comprehensive but concise."""
                    },
                    {
                        "role": "user",
                        "content": f"Please analyze this document and provide a comprehensive summary in the same language as the document:\n\n{text[:4000]}..."
                    }
                ],
                max_tokens=500,
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating document summary: {str(e)}"

    def chat_with_document(self, question: str):
        """Chat with the document content using conversational retrieval"""
        if self.document_conversation_chain is None:
            return "No document processed yet. Please process a document first."

        try:
            response = self.document_conversation_chain({"question": question})
            return response['answer']
        except Exception as e:
            return f"Error in document conversation: {str(e)}"

    def reset_document_conversation(self):
        """Reset the document conversation memory"""
        if self.document_memory is not None:
            self.document_memory.clear()
            return "Document conversation history cleared!"
        return "No document conversation to reset."

    def search_document_vector_store(self, query: str, k: int = 3):
        """Search the document vector store for relevant chunks"""
        if self.document_vector_store is None:
            return ["No document processed yet. Process a document first."]

        try:
            docs = self.document_vector_store.similarity_search(query, k=k)
            return [doc.page_content for doc in docs]
        except Exception as e:
            return [f"Error searching document: {str(e)}"]